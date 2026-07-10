"""
extract.py
==========
Walk the raw MOOC folders and turn every .docx / .pdf / .vtt file into clean text.

Design notes (why it looks like this):
  * The dataset is Google-Docs exports. PDFs come out with weird wide spacing
    ("C o m p u t e r   S e c u r i t y") and doubled spaces -> we normalise.
  * .docx files hold the real teaching content (articles). We read paragraphs
    AND tables.
  * .vtt are subtitle files -> we strip timestamps/cue headers, keep spoken text.
  * We SKIP image folders and binary junk automatically (only 3 extensions read).
"""
from __future__ import annotations
import re
from pathlib import Path
from typing import Iterator

import docx  # python-docx
from pdfminer.high_level import extract_text as pdf_extract_text

# Extensions we actually read. Everything else (jpg, png, xlsx) is ignored.
TEXT_EXTS = {".docx", ".pdf", ".vtt"}


# ----------------------------------------------------------------------
# Cleaning helpers
# ----------------------------------------------------------------------
def _collapse_spaced_out_text(text: str) -> str:
    """
    Google-Docs PDF exports often space every character:
        'C o m p u t e r' -> 'Computer'
    Heuristic: if a line has an unusually high ratio of single-char 'words',
    remove the spurious spaces between single characters.
    """
    lines = []
    for line in text.split("\n"):
        tokens = line.split(" ")
        singles = sum(1 for t in tokens if len(t) == 1)
        if tokens and singles / max(len(tokens), 1) > 0.5 and len(tokens) > 4:
            # Rejoin: collapse runs of single chars, keep real word gaps (double space)
            line = re.sub(r"(?<=\w) (?=\w)", "", line)   # remove single-char gaps
        lines.append(line)
    return "\n".join(lines)


def clean_text(text: str) -> str:
    """Normalise whitespace and strip export artefacts. Safe for all sources."""
    if not text:
        return ""
    text = text.replace("\u00a0", " ")           # non-breaking space
    text = _collapse_spaced_out_text(text)
    text = re.sub(r"[ \t]+", " ", text)          # collapse runs of spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)       # cap blank lines
    # Drop obvious markdown-ish artefact markers left by the export
    text = re.sub(r"\[Alt:.*?\]", " ", text, flags=re.DOTALL)
    text = re.sub(r"\[Image.*?\]", " ", text, flags=re.DOTALL)
    return text.strip()


# ----------------------------------------------------------------------
# Per-format readers
# ----------------------------------------------------------------------
def read_docx(path: Path) -> str:
    """Paragraphs + table cells from a .docx file."""
    try:
        d = docx.Document(str(path))
    except Exception as e:                      # corrupt / not a real docx
        print(f"  [warn] could not open docx {path.name}: {e}")
        return ""
    parts: list[str] = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    try:
        return pdf_extract_text(str(path)) or ""
    except Exception as e:
        print(f"  [warn] could not read pdf {path.name}: {e}")
        return ""


def read_vtt(path: Path) -> str:
    """Strip WEBVTT header, timestamps and cue settings; keep spoken lines."""
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        print(f"  [warn] could not read vtt {path.name}: {e}")
        return ""
    out: list[str] = []
    for line in raw.splitlines():
        s = line.strip()
        if not s or s == "WEBVTT":
            continue
        if "-->" in s:                          # timestamp cue line
            continue
        if s.isdigit():                         # numeric cue index
            continue
        s = re.sub(r"<[^>]+>", "", s)           # inline tags <c> etc.
        out.append(s)
    # de-duplicate consecutive identical lines (common in captions)
    deduped = []
    for line in out:
        if not deduped or deduped[-1] != line:
            deduped.append(line)
    return " ".join(deduped)


_READERS = {".docx": read_docx, ".pdf": read_pdf, ".vtt": read_vtt}


# ----------------------------------------------------------------------
# Directory walk
# ----------------------------------------------------------------------
def iter_documents(raw_dir: Path) -> Iterator[dict]:
    """
    Yield one dict per readable file:
        {source, module, topic, text}
    'module' = top-level folder (e.g. 'MOOC 1'); 'topic' = second-level if present.
    """
    raw_dir = Path(raw_dir)
    files = sorted(p for p in raw_dir.rglob("*") if p.suffix.lower() in TEXT_EXTS)
    print(f"Found {len(files)} readable files under {raw_dir}")
    for path in files:
        rel = path.relative_to(raw_dir)
        parts = rel.parts
        module = parts[0] if len(parts) > 0 else "unknown"
        topic  = parts[1] if len(parts) > 1 else module
        reader = _READERS[path.suffix.lower()]
        text = clean_text(reader(path))
        if len(text) < 20:        # nothing useful extracted
            continue
        yield {
            "source": str(rel).replace("\\", "/"),
            "module": module,
            "topic": topic,
            "text": text,
        }


if __name__ == "__main__":
    # Quick smoke test: python -m src.extract
    from src.config import RAW_DIR
    n = 0
    total_chars = 0
    for doc in iter_documents(RAW_DIR):
        n += 1
        total_chars += len(doc["text"])
        if n <= 3:
            print(f"\n--- {doc['source']} ({len(doc['text'])} chars) ---")
            print(doc["text"][:250])
    print(f"\nExtracted {n} documents, {total_chars:,} characters total.")
